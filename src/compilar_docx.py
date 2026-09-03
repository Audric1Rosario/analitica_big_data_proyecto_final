"""
Compilador de Informe Final en Formato Microsoft Word (.docx)
Proyecto Final: Aplicaciones Analíticas de Big Data (UAPA)
Equipo: Audric Rosario & Orlando Benítez
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn


def set_cell_background(cell, fill_hex):
    """Aplica color de fondo a una celda de tabla."""
    tc_pr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Ajusta el espaciado interno de una celda."""
    tc_pr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tc_pr.append(tcMar)


def crear_documento_word(md_path: str = "INFORME_FINAL_CLARO.md", docx_out: str = "INFORME_FINAL_CLARO.docx"):
    print(f"[DOCX] Leyendo informe en Markdown desde: {md_path}...")
    with open(md_path, "r", encoding="utf-8") as f:
        lineas = f.readlines()

    doc = Document()

    # Configuración de márgenes estándar (1 pulgada)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Estilos de encabezado y párrafos
    color_primario = RGBColor(218, 41, 28)      # Rojo Claro corporativo (#DA291C)
    color_secundario = RGBColor(15, 23, 42)    # Azul marino oscuro (#0F172A)
    color_texto = RGBColor(30, 41, 59)          # Gris pizarra oscuro

    # Configurar estilo Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = color_texto
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    # Portada Especial
    p_portada_uni = doc.add_paragraph()
    p_portada_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_uni = p_portada_uni.add_run("UNIVERSIDAD ABIERTA PARA ADULTOS (UAPA)\nVICERRECTORÍA DE INVESTIGACIÓN Y POSGRADO\nMAESTRÍA EN ANALÍTICA DE BIG DATA E INTELIGENCIA DE NEGOCIOS\n")
    run_uni.bold = True
    run_uni.font.size = Pt(13)
    run_uni.font.color.rgb = color_secundario

    doc.add_paragraph()  # Espacio

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tit = p_titulo.add_run("PROYECTO FINAL:\nAUDITORÍA DE EXPERIENCIA DEL CLIENTE Y SENTIMIENTO DE MARCA EN TELECOMUNICACIONES MEDIANTE NLP Y YOUTUBE DATA API V3")
    run_tit.bold = True
    run_tit.font.size = Pt(16)
    run_tit.font.color.rgb = color_primario

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Caso de Estudio: Claro República Dominicana (@clarord)\n")
    run_sub.italic = True
    run_sub.font.size = Pt(13)

    doc.add_paragraph()  # Espacio

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run("ASIGNATURA: APLICACIONES ANALÍTICAS DE BIG DATA\nFACILITADOR: PROFESOR DE LA ASIGNATURA\n\nINTEGRANTES DEL EQUIPO:\nAUDRIC ROSARIO\nORLANDO BENÍTEZ\n\nSANTIAGO / SANTO DOMINGO, REPÚBLICA DOMINICANA\nSEPTIEMBRE 2026")
    r_meta.bold = True
    r_meta.font.size = Pt(11)
    r_meta.font.color.rgb = color_secundario

    doc.add_page_break()

    # Procesar las secciones del Markdown
    i = 0
    tabla_buffer = []

    while i < len(lineas):
        line = lineas[i].strip()

        # Detectar Tablas en Markdown
        if line.startswith("|") and line.endswith("|"):
            tabla_buffer.append(line)
            i += 1
            continue
        elif tabla_buffer:
            # Renderizar tabla acumulada
            filas_datos = [f for f in tabla_buffer if not re.match(r"^\|[\s\-:]+\|$", f)]
            if len(filas_datos) >= 1:
                cols = [c.strip() for c in filas_datos[0].split("|")[1:-1]]
                table = doc.add_table(rows=len(filas_datos), cols=len(cols))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for row_idx, fila_str in enumerate(filas_datos):
                    celdas_texto = [c.strip() for c in fila_str.split("|")[1:-1]]
                    row = table.rows[row_idx]
                    for col_idx, txt in enumerate(celdas_texto):
                        if col_idx < len(row.cells):
                            cell = row.cells[col_idx]
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            cell.text = txt.replace("**", "")
                            # Estilizar encabezado
                            if row_idx == 0:
                                set_cell_background(cell, "DA291C")  # Rojo Claro
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        r.font.bold = True
                                        r.font.color.rgb = RGBColor(255, 255, 255)
                                        r.font.size = Pt(9.5)
                            else:
                                if row_idx % 2 == 1:
                                    set_cell_background(cell, "F8FAFC")
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        r.font.size = Pt(9.5)
                                        r.font.color.rgb = color_texto
                            set_cell_margins(cell, 80, 80, 100, 100)

                doc.add_paragraph()  # Espacio post tabla
            tabla_buffer = []

        # Títulos de Nivel 2 (Secciones 6.1 a 6.21)
        if line.startswith("## 6."):
            h2 = doc.add_paragraph()
            h2.paragraph_format.space_before = Pt(14)
            h2.paragraph_format.space_after = Pt(6)
            h2.paragraph_format.keep_with_next = True
            run_h2 = h2.add_run(line.replace("##", "").strip())
            run_h2.bold = True
            run_h2.font.size = Pt(13.5)
            run_h2.font.color.rgb = color_primario

        # Títulos de Nivel 3 (Subsecciones)
        elif line.startswith("### "):
            h3 = doc.add_paragraph()
            h3.paragraph_format.space_before = Pt(10)
            h3.paragraph_format.space_after = Pt(4)
            h3.paragraph_format.keep_with_next = True
            run_h3 = h3.add_run(line.replace("###", "").strip())
            run_h3.bold = True
            run_h3.font.size = Pt(12)
            run_h3.font.color.rgb = color_secundario

        # Viñetas
        elif line.startswith("* ") or line.startswith("- "):
            p_bullet = doc.add_paragraph(style='List Bullet')
            p_bullet.paragraph_format.space_after = Pt(3)
            texto_bullet = line[2:].strip()
            # Tratar negritas
            partes = re.split(r'(\*\*.*?\*\*)', texto_bullet)
            for parte in partes:
                if parte.startswith("**") and parte.endswith("**"):
                    r = p_bullet.add_run(parte[2:-2])
                    r.bold = True
                else:
                    p_bullet.add_run(parte)

        # Listas numeradas
        elif re.match(r"^\d+\.\s", line):
            p_num = doc.add_paragraph(style='List Number')
            p_num.paragraph_format.space_after = Pt(3)
            texto_num = re.sub(r"^\d+\.\s", "", line)
            partes = re.split(r'(\*\*.*?\*\*)', texto_num)
            for parte in partes:
                if parte.startswith("**") and parte.endswith("**"):
                    r = p_num.add_run(parte[2:-2])
                    r.bold = True
                else:
                    p_num.add_run(parte)

        # Párrafos normales
        elif line:
            # Omitir separadores horizontales y bloques de diagramas mermaid en texto plano
            if line.startswith("---") or line.startswith("```"):
                i += 1
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            partes = re.split(r'(\*\*.*?\*\*)', line)
            for parte in partes:
                if parte.startswith("**") and parte.endswith("**"):
                    r = p.add_run(parte[2:-2])
                    r.bold = True
                else:
                    p.add_run(parte)

        i += 1

    doc.save(docx_out)
    print(f"[EXITO] Documento Word generado con éxito en: {docx_out}")
    return docx_out


if __name__ == "__main__":
    crear_documento_word()
